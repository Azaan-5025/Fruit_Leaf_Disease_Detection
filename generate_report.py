"""
Generate Project Report with user-provided text content exactly as requested.
Keeps formatting, images, tables, code blocks intact.
"""

import os, json
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches as DI, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTFILE = "Project_Report_v3.docx"
EDA_DIR = "eda_outputs"
RESULTS_DIR = "results"

def load_model_data():
    p = "models/comparison_results.json"
    if os.path.exists(p):
        with open(p) as f: return json.load(f)
    return []

MODEL_DATA = load_model_data()

def ensure_dirs():
    os.makedirs(EDA_DIR, exist_ok=True)
    os.makedirs(RESULTS_DIR, exist_ok=True)
    # Preprocessing flow
    if not os.path.exists(f"{EDA_DIR}/preprocessing_flow.png"):
        fig,ax=plt.subplots(figsize=(12,3)); ax.axis('off')
        steps=[("Load Image\n(RGB)","#E91E63"),("Resize\n224×224","#9C27B0"),
               ("Gaussian Blur\n5×5 kernel","#3F51B5"),("Normalize\n[0,1]","#00BCD4"),
               ("Sakaguchi Tensors\n8×8, 12×12, 16×16","#4CAF50")]
        for i,(lbl,clr) in enumerate(steps):
            x=0.05+i*0.19
            rect=plt.Rectangle((x,0.2),0.16,0.6,facecolor=clr,edgecolor='white',lw=2,alpha=0.9,transform=ax.transAxes)
            ax.add_patch(rect)
            ax.text(x+0.08,0.5,lbl,transform=ax.transAxes,ha='center',va='center',color='white',fontsize=9,fontweight='bold')
            if i<4: ax.annotate('',xy=(x+0.18,0.5),xytext=(x+0.16,0.5),xycoords='axes fraction',
                                textcoords='axes fraction',arrowprops=dict(arrowstyle='->',color='#333',lw=2))
        ax.set_title("Sakaguchi Preprocessing Pipeline",fontsize=13,fontweight='bold',pad=10)
        plt.tight_layout(); fig.savefig(f"{EDA_DIR}/preprocessing_flow.png",dpi=150); plt.close()

def build_report():
    doc = Document()
    for s in doc.sections:
        s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5); s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)

    def heading(text, level=1):
        h=doc.add_heading(text,level=level)
        for r in h.runs: r.font.color.rgb=RGBColor(0x0D,0x47,0xA1)

    def para(text, bold=False, italic=False, sz=11):
        p=doc.add_paragraph()
        r=p.add_run(text); r.font.size=Pt(sz); r.font.name='Calibri'; r.bold=bold; r.italic=italic
        p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.15

    def bullet(text, sz=11):
        p=doc.add_paragraph(style='List Bullet')
        p.clear(); r=p.add_run(text); r.font.size=Pt(sz); r.font.name='Calibri'
        p.paragraph_format.space_after=Pt(3)

    def code_block(lines_text):
        for line in lines_text.strip().split('\n'):
            p=doc.add_paragraph()
            r=p.add_run(line); r.font.name='Consolas'; r.font.size=Pt(8)
            r.font.color.rgb=RGBColor(0x33,0x33,0x33)
            p.paragraph_format.space_after=Pt(0); p.paragraph_format.space_before=Pt(0)
            p.paragraph_format.line_spacing=1.0; p.paragraph_format.left_indent=Cm(0.5)

    def add_img(path, width=5.5, caption=None):
        if os.path.exists(path):
            doc.add_picture(path, width=DI(width))
            doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
            if caption:
                cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                r=cp.add_run(caption); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=RGBColor(0x66,0x66,0x66)

    def add_table(headers, rows):
        t=doc.add_table(rows=len(rows)+1,cols=len(headers)); t.style='Light Grid Accent 1'
        t.alignment=WD_TABLE_ALIGNMENT.CENTER
        for i,h in enumerate(headers):
            c=t.rows[0].cells[i]; c.text=h
            for p in c.paragraphs:
                for r in p.runs: r.bold=True
        for ri,row in enumerate(rows):
            for ci,val in enumerate(row): t.rows[ri+1].cells[ci].text=str(val)
        doc.add_paragraph("")

    # =========================================================
    # TITLE PAGE
    # =========================================================
    for _ in range(4): doc.add_paragraph("")
    tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=tp.add_run("Vellore Institute of Technology"); r.font.size=Pt(16); r.bold=True; r.font.color.rgb=RGBColor(0x0D,0x47,0xA1)
    tp2=doc.add_paragraph(); tp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=tp2.add_run("School of Computer Science and Engineering"); r2.font.size=Pt(13); r2.font.color.rgb=RGBColor(0x33,0x33,0x33)
    doc.add_paragraph(""); doc.add_paragraph("")
    tp3=doc.add_paragraph(); tp3.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r3=tp3.add_run("Multi Input Fruit and Leaf Disease Detection\nUsing Deep Learning"); r3.font.size=Pt(26); r3.bold=True; r3.font.color.rgb=RGBColor(0x0D,0x47,0xA1)
    doc.add_paragraph("")
    tp4=doc.add_paragraph(); tp4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r4=tp4.add_run("DS Project Report"); r4.font.size=Pt(16); r4.font.color.rgb=RGBColor(0x55,0x55,0x55)
    for _ in range(3): doc.add_paragraph("")
    info=doc.add_paragraph(); info.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for line in ["Submitted by Syed Azaan Hussain","Register No 24BCE5025","","Under the guidance of Pattabhiraman V","","Vellore Institute of Technology"]:
        rn=info.add_run(line+"\n"); rn.font.size=Pt(12); rn.font.name='Calibri'
        if "Submitted" in line or "guidance" in line: rn.bold=True
    doc.add_page_break()

    # =========================================================
    # TABLE OF CONTENTS
    # =========================================================
    heading("Table of Contents", 1)
    for item in ["1 Abstract","2 Introduction","3 Literature Survey","4 Dataset and Description",
                 "5 Exploratory Data Analysis EDA Code and Visuals","6 Data Pre Processing Code and Output",
                 "7 Model Implementations Code with Input Output","8 Prediction Results with Accuracy Table and Graphs",
                 "9 Comparison Proposed Work vs Existing Work","10 Results and Discussion",
                 "11 Conclusions with Future Enhancement","12 References"]:
        p=doc.add_paragraph(item); p.paragraph_format.space_after=Pt(2)
    doc.add_page_break()

    # =========================================================
    # 1 ABSTRACT
    # =========================================================
    heading("1 Abstract", 1)
    para("Crop diseases mess up farming everywhere and I think they cause losses around 30 percent in some areas from what I read. For small farmers in places like developing countries its tough because most rely on that and getting experts to look at leaves or fruits by eye is expensive and takes too long not practical at all.")
    para("In my project I put together a system that diagnoses diseases on fruits and vegetable leaves using computer vision basically deep learning. The key thing I tried was processing images into these multiple sizes like small ones 8 by 8 pixels 12 by 12 16 by 16 and the standard bigger 224 by 224. It was inspired by something called the Sakaguchi reaction I guess to capture different levels of detail from the same picture.")
    para("I ran these through four different models a CNN using MobileNetV2 base an ANN with some layers ResNet50 that I tweaked a little and then an SVM that took features from the MobileNetV2. The dataset was pretty big about 82 thousand images spread over 30 categories from eight different plants. On validation the CNN hit 85.05 percent accuracy and the SVM got 82.34 percent which seemed okay. I built it into a Streamlit app so you can upload an image and see predictions fast.")
    para("This might actually help farmers out there but im not totally sure about real field conditions yet it feels promising though.")
    doc.add_page_break()

    # =========================================================
    # 2 INTRODUCTION
    # =========================================================
    heading("2 Introduction", 1)
    para("Food comes from plants thats obvious but diseases from fungi bacteria or whatever wipe out about a third of crops yearly. You notice spots wilting discoloration on leaves and fruits stuff like that. Back in the day farmers would call in experts but access is limited and even experts might disagree its subjective.")
    para("These days with smartphone cameras being so good we can use deep learning to classify plant images on the spot. CNNs pick up on patterns like colors textures shapes really well in photos.")
    para("What drew me to this was seeing how most projects stick to one clean dataset like PlantVillage all lab perfect images. But in reality lighting is bad angles are off backgrounds messy. Plus not many compare various models side by side on the same mixed data.")
    para("I combined PlantVillage with a fruits dataset from Kaggle ending up with 30 classes covering leaves and whole fruits. Added that multi size preprocessing step and trained four models consistently. The report covers literature next then datasets eda preprocessing models code results comparison discussion and some wrap up with future thoughts.")
    para("Combining datasets probably makes it trickier I wonder sometimes but it seems needed for something practical.")
    doc.add_page_break()

    # =========================================================
    # 3 LITERATURE SURVEY
    # =========================================================
    heading("3 Literature Survey", 1)
    para("People have been working on deep learning for plant diseases for years now. I went through some main papers that influenced my approach.")
    para("Back in 2016 Mohanty and his team used AlexNet and GoogLeNet on the PlantVillage dataset 54 thousand images 38 classes. With transfer learning they reached 99.35 percent accuracy but it tanked when tested on real world non lab photos.")
    para("In 2019 Too and others compared VGG16 InceptionV4 ResNets DenseNet on that same data. DenseNet121 came out on top at 99.75 percent thanks to its feature reuse I think.")
    para("Brahimi in 2017 looked at tomatoes specifically using GoogleNet and AlexNet for nine classes fine tuned to 99.18 percent.")
    para("Ferentinos tried five CNNs on PlantVillage in 2018 best was 99.53 percent and they broke down per class where similar diseases confused the models.")
    para("All that work stays with one pristine dataset though. In my case I mixed two datasets added the multi resolution input and tested both deep models and a non deep one like SVM. The accuracies might not be as high but it feels more like real life to me.")
    doc.add_page_break()

    # =========================================================
    # 4 DATASET AND DESCRIPTION
    # =========================================================
    heading("4 Dataset and Description", 1)

    heading("4.1 PlantVillage Dataset", 2)
    para("For the dataset I started with PlantVillage from Kaggle over 50 thousand leaf images healthy and diseased. I selected 15 classes from tomato pepper potato. Tomato covers 10 issues bacterial spot early late blight leaf mold septoria spot spider mites target spot yellow leaf curl virus mosaic virus and healthy. Pepper has bacterial spot healthy potato early late blight healthy.")
    para("Source https www.kaggle.com datasets emmarex plantdisease")

    heading("4.2 Fruits Disease Dataset", 2)
    para("Then the fruits disease dataset includes apples bananas grapes mangos oranges each with fresh rotten formalin mixed states so 15 classes. Images show whole fruits from various angles different backgrounds.")
    para("Source https www.kaggle.com datasets sriramr fruits diseases")

    heading("4.3 Consolidated Dataset", 2)
    para("I merged everything using a script into one folder 30 classes total around 82 thousand images.")

    add_table(
        ["Parameter", "Value"],
        [["Total Classes", "30 (15 from plantvillage 15 fruits)"],
         ["Total Images", "about 82000"],
         ["Image Input Size", "224 x 224 x 3 rgb"],
         ["Train Validation Split", "80 20"],
         ["Batch Size", "32"]]
    )
    para("Figure 4.1 is the dataset stats summary.")
    add_img(f"{EDA_DIR}/dataset_statistics.png", 5.0, "Figure 4.1 Dataset Statistics Summary")

    heading("4.4 Data Consolidation Code", 2)
    para("To combine the data I wrote this script", bold=True)
    code_block('''import os
import shutil

def consolidate():
    target_root = r"c:\\Users\\Syed Azaan Hussain\\DS_Project\\data\\training_data"
    
    pv_source = r"c:\\Users\\Syed Azaan Hussain\\DS_Project\\data\\PlantVillage"
    pv_classes = [
        "Pepper__bell___Bacterial_spot", "Pepper__bell___healthy", "Potato___Early_blight",
        "Potato___Late_blight", "Potato___healthy", "Tomato_Bacterial_spot", 
        "Tomato_Early_blight", "Tomato_Late_blight", "Tomato_Leaf_Mold", 
        "Tomato_Septoria_leaf_spot", "Tomato_Spider_mites_Two_spotted_spider_mite", 
        "Tomato__Target_Spot", "Tomato__Tomato_YellowLeaf__Curl_Virus", 
        "Tomato__Tomato_mosaic_virus", "Tomato_healthy"
    ]

    fruit_source = r"c:\\Users\\Syed Azaan Hussain\\DS_Project\\data\\fruits Dataset\\train"
    fruit_classes = ["Apple", "Banana", "Grape", "Mango", "Orange"]

    if not os.path.exists(target_root):
        os.makedirs(target_root)

    for cls in pv_classes:
        src_path = os.path.join(pv_source, cls)
        if os.path.exists(src_path):
            dst_path = os.path.join(target_root, cls)
            print(f"Copying PV class {cls}")
            if os.path.exists(dst_path): shutil.rmtree(dst_path)
            shutil.copytree(src_path, dst_path)

    for fruit in fruit_classes:
        fruit_path = os.path.join(fruit_source, fruit)
        if os.path.exists(fruit_path):
            for condition in os.listdir(fruit_path):
                cond_path = os.path.join(fruit_path, condition)
                if os.path.isdir(cond_path):
                    new_cls_name = f"{fruit}___{condition}"
                    dst_path = os.path.join(target_root, new_cls_name)
                    print(f"Copying Fruit class {new_cls_name}")
                    if os.path.exists(dst_path): shutil.rmtree(dst_path)
                    shutil.copytree(cond_path, dst_path)

if __name__ == "__main__":
    consolidate()''')
    para("It ran fine but the paths are just for my setup kind of specific.")
    doc.add_page_break()

    # =========================================================
    # 5 EDA
    # =========================================================
    heading("5 Exploratory Data Analysis EDA Code and Visuals", 1)
    para("Before diving into models I did some exploratory data analysis to check balances and visuals using python and matplotlib.")
    para("The eda code starts with", bold=True)
    code_block('''import os
import matplotlib.pyplot as plt

data_dir = "data/training_data"
classes = sorted(os.listdir(data_dir))
print(f"Total classes found: {len(classes)}")

counts = {}
for cls in classes:
    path = os.path.join(data_dir, cls)
    if os.path.isdir(path):
        counts[cls] = len(os.listdir(path))

print(f"Total images: {sum(counts.values())}")

categories = {}
for cls_name in classes:
    for key in ["Tomato","Pepper","Potato","Apple","Banana","Grape","Mango","Orange"]:
        if key in cls_name:
            categories[key] = categories.get(key, 0) + 1
            break

plt.figure(figsize=(10, 5))
plt.barh(list(categories.keys()), list(categories.values()))
plt.xlabel("Number of Classes")
plt.title("Classes per Plant Fruit Category")
plt.tight_layout()
plt.savefig("eda_outputs/class_distribution.png")
plt.show()

healthy = sum(1 for c in classes if "healthy" in c.lower() or "Fresh" in c)
diseased = len(classes) - healthy
plt.figure(figsize=(5, 5))
plt.pie([healthy, diseased], labels=["Healthy", "Diseased"],
        autopct="%1.0f%%", colors=["#4CAF50", "#F44336"])
plt.title("Healthy vs Diseased Classes")
plt.savefig("eda_outputs/healthy_vs_diseased.png")
plt.show()''')

    para("I think that covers the basics of looking at class distributions and maybe some sample images but it got a bit messy when I plotted everything. Some classes had way more images than others which might affect training I suppose.")
    para("I started by setting up the data directory with the training data. There are classes listed in there, and I sorted them to get a clean list. It turned out there were quite a few, like 30 in total. I printed that out to check.")
    para("For each class folder, I counted the images inside. Some paths were directories, so that worked. The total images added up to around 50,000 or so, I think, but I just summed the counts to confirm.")
    para("Then I tried grouping them into categories, like Tomato, Pepper, Potato, and fruits such as Apple, Banana, Grape, Mango, Orange. I looped through the class names and matched keywords. If a class had Tomato in it, it went under that key, incrementing the count. Not everything fit perfectly, but it gave a rough idea.")
    para("I made a horizontal bar plot for the categories, with sizes 10 by 5. Labeled it Number of Classes on the x, and title Classes per Plant Fruit Category. Saved it as class_distribution.png in eda_outputs, and showed it.")

    add_img(f"{EDA_DIR}/class_distribution.png", 5.5, "Figure 5.1 Number of classes per plant fruit category")
    para("Tomato ended up with the most, 10 classes out of 30. Fruits like Apple had 3 each, or something close.")

    para("Next, I checked healthy versus diseased. I summed up classes with healthy or Fresh in the lowercase name. That gave about 10 healthy ones. Diseased was the rest, so 20. Made a pie chart, 5 by 5, with labels Healthy and Diseased, percentages to one decimal, colors green and red. Title Healthy vs Diseased Classes, saved as healthy_vs_diseased.png.")
    
    add_img(f"{EDA_DIR}/healthy_vs_diseased.png", 3.5, "Figure 5.2 Proportion of healthy vs diseased classes")
    para("About 33 percent healthy, which seems low, but maybe the dataset focuses on diseases.")

    if os.path.exists(f"{EDA_DIR}/class_list.png"):
        add_img(f"{EDA_DIR}/class_list.png", 5.5, "Figure 5.3 Complete list of 30 disease classes")
        
    para("I also listed all 30 classes, but didnt plot that formally, just noted it.")
    para("Overall, Tomato dominates everything, 10 classes. Fruits are minor, each around 3. The dataset feels somewhat balanced, not super skewed, though some classes probably have fewer images. I didnt count per class yet, that might be next.")
    doc.add_page_break()

    # =========================================================
    # 6 DATA PRE-PROCESSING
    # =========================================================
    heading("6 Data Pre Processing Code and Output", 1)

    para("Moving to preprocessing, I followed this Sakaguchi pipeline to clean images and handle multiple sizes. It standardizes them, smooths out noise with Gaussian blur, and resizes for different resolutions to catch textures better.")
    para("I think the multi sizes are useful for details at coarse, medium, fine levels, though im not totally sure if it makes a huge difference.")
    
    add_img(f"{EDA_DIR}/preprocessing_flow.png", 6.0, "Figure 6.1 Preprocessing Pipeline Flow")

    para("The code is in src/preprocessing.py. It imports cv2 and numpy.", bold=True)
    code_block('''import cv2
import numpy as np

def load_image(image_file, target_size=(224, 224)):
    if hasattr(image_file, 'read'):
        file_bytes = np.asarray(bytearray(image_file.read()), dtype=np.uint8)
        image = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
    else:
        image = cv2.imread(image_file)
    
    if image is None:
        raise ValueError("Could not load image")
        
    image = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
    return image

def apply_noise_smoothing(image, kernel_size=(5, 5)):
    return cv2.GaussianBlur(image, kernel_size, 0)

def normalize_image(image):
    return image.astype(np.float32) / 255.0

def resize_image(image, size):
    return cv2.resize(image, size)

def sakaguchi_tensor_conversion(image):
    tensors = {}
    sizes = [(8, 8), (12, 12), (16, 16)]
    
    for size in sizes:
        resized = resize_image(image, size)
        tensors[f"{size[0]}x{size[1]}"] = resized
        
    return tensors

def preprocess_pipeline(image_file, target_size=(224, 224)):
    image = load_image(image_file)
    image = apply_noise_smoothing(image)
    image = resize_image(image, target_size)
    image = normalize_image(image)
    tensors = sakaguchi_tensor_conversion(image)
    return image, tensors''')

    para("Load image function takes a file, target size default 224 by 224. If its a file object, it reads bytes and decodes with imdecode, color mode. Otherwise, just imread. Converts BGR to RGB, returns the image. If none, raises value error.")
    para("Then apply noise smoothing, Gaussian blur with kernel 5 by 5, sigma zero.")
    para("Normalize divides by 255 to float32, 0 to 1 range.")
    para("Resize with cv2 resize to the size.")
    para("Sakaguchi conversion makes tensors dictionary, sizes 8x8, 12x12, 16x16. Resizes the original image to each, stores under keys like 8x8.")
    para("The pipeline loads, smooths, resizes to target, normalizes, then makes the tensors from the normalized one, I suppose. Returns image and tensors.")
    para("My notes cut off, but thats the main flow.")

    para("For augmentation during training, I used Keras ImageDataGenerator. Rescale 1/255, preprocessing function is the sakaguchi one, rotation 20 degrees, zoom 0.15, shifts 0.2 width and height, shear 0.15, horizontal flip true, validation split 0.2, fill nearest.", bold=True)
    code_block('''datagen = ImageDataGenerator(
    rescale=1./255,
    preprocessing_function=sakaguchi_preprocessor,
    rotation_range=20,
    zoom_range=0.15,
    width_shift_range=0.2,
    height_shift_range=0.2,
    shear_range=0.15,
    horizontal_flip=True,
    validation_split=0.2,
    fill_mode="nearest"
)''')
    para("That adds variety, I guess, to avoid overfitting.")

    para("Since models need multiple inputs, I wrapped the generator in a MultiInputGeneratorWrapper, subclass of tf.keras.utils.Sequence. Init takes the generator.", bold=True)
    code_block('''class MultiInputGeneratorWrapper(tf.keras.utils.Sequence):
    def __init__(self, generator):
        self.generator = generator

    def __getitem__(self, index):
        img_batch, labels = self.generator[index]
        t8_batch, t12_batch, t16_batch = [], [], []
        for img in img_batch:
            tensors = sakaguchi_tensor_conversion(img)
            t8_batch.append(tensors['8x8'])
            t12_batch.append(tensors['12x12'])
            t16_batch.append(tensors['16x16'])
        return {
            "img_input": np.array(img_batch),
            "t8_input": np.array(t8_batch),
            "t12_input": np.array(t12_batch),
            "t16_input": np.array(t16_batch)
        }, labels''')
    para("Getitem gets batch of images and labels from generator. Then for each image in batch, convert to tensors on the fly, append the 8x8, 12x12, 16x16 to lists. Return numpy arrays of img batch, t8, t12, t16, and labels.")
    para("Storing everything would take too much space, so on the fly makes sense.")

    para("The preprocessing stages: load gives variable raw RGB. Resize to 224x224x3 standard. Blur same shape, smoothed. Normalize still 224x224x3 but 0-1. Then Sakaguchi gives 8x8x3 coarse, 12x12x3 medium, 16x16x3 fine.")
    para("Later they get flattened into features, I think.")

    add_table(
        ["Stage", "Output Shape", "Description"],
        [["Load Image", "Variable", "Raw RGB"],
         ["Resize", "224 x 224 x 3", "Standard size"],
         ["Gaussian Blur", "224 x 224 x 3", "Smoothed"],
         ["Normalize", "224 x 224 x 3", "To 0-1"],
         ["Sakaguchi 8x8", "8 x 8 x 3", "Coarse"],
         ["Sakaguchi 12x12", "12 x 12 x 3", "Medium"],
         ["Sakaguchi 16x16", "16 x 16 x 3", "Fine"]]
    )
    doc.add_page_break()

    # =========================================================
    # 7 MODEL IMPLEMENTATIONS
    # =========================================================
    heading("7 Model Implementations Code with Input Output", 1)
    para("For models, I implemented four, all multi input. First, the common multi input layers. Function get_multi_input_layers, shape 224x224x3. Inputs for img, t8 8x8x3, t12 12x12x3, t16 16x16x3, named accordingly.")

    para("That sets up the base, then each model builds from there. Some might combine them differently, but I tried variations.", bold=True)
    code_block('''def get_multi_input_layers(img_shape=(224, 224, 3)):
    img_input = Input(shape=img_shape, name="img_input")
    t8_input = Input(shape=(8, 8, 3), name="t8_input")
    t12_input = Input(shape=(12, 12, 3), name="t12_input")
    t16_input = Input(shape=(16, 16, 3), name="t16_input")

    t8_feat = Dense(32, activation='relu')(Flatten()(t8_input))
    t12_feat = Dense(64, activation='relu')(Flatten()(t12_input))
    t16_feat = Dense(128, activation='relu')(Flatten()(t16_input))

    return [img_input, t8_input, t12_input, t16_input], [t8_feat, t12_feat, t16_feat]''')
    para("I started with this get multi input layers function because it handles the main image plus these smaller tensor inputs at different scales like 8 by 8, 12 by 12, and 16 by 16. It takes the big 224 by 224 by 3 image input and then creates these auxiliary inputs from it using average pooling. For each small one, it applies pooling to get down to that size, flattens it a bit, but actually in the code its just the pooled output as input, and then later dense layers pull features out. The function returns all the inputs and the features from dense layers with relu activation, 32 units for the 8, 64 for 12, 128 for 16. So overall, inputs four tensors but outputs the features from the small ones too, which get concatenated later.")

    para("That setup feeds into the models I built. For the first one, the CNN with MobileNetV2, I used the pre trained version from imagenet, froze the base so it wouldnt train, and added the auxiliary features. The base takes the main image input, global average pooling on its output, then concatenate with the aux feats, dense 256 relu, dropout 0.3, and softmax for 30 classes. Compiled with adam and categorical crossentropy, accuracy metric. Its lightweight, which I think makes it good for an app or something mobile.", bold=True)
    code_block('''def build_cnn_model(num_classes=30, input_shape=(224, 224, 3)):
    inputs, aux_feats = get_multi_input_layers(input_shape)
    base_model = MobileNetV2(weights='imagenet', include_top=False, input_tensor=inputs[0])
    base_model.trainable = False
    x = GlobalAveragePooling2D()(base_model.output)
    combined = Concatenate()([x] + aux_feats)  # 1280 + 32 + 64 + 128 = 1504 features
    x = Dense(256, activation='relu')(combined)
    x = Dropout(0.3)(x)
    predictions = Dense(num_classes, activation='softmax')(x)
    model = Model(inputs=inputs, outputs=predictions)
    model.compile(optimizer='adam', loss='categorical_crossentropy', metrics=['accuracy'])
    return model''')

    para("Then there is the ANN MLP model, where I flattened the big image right away, which is 224 by 224 by 3 so huge input, dense 1024 relu, batch norm, dropout 0.5, another dense 512 relu, batch norm again, then concatenate the aux feats, more dense layers down to 128 relu, then predictions. Used adam with lower lr 0.0005. This one did not do well, like I expected because flattening loses all the spatial info, no convolutions to help.", bold=True)
    code_block('''def build_ann_model(num_classes=30, input_shape=(224, 224, 3)):
    inputs, aux_feats = get_multi_input_layers(input_shape)
    x = Flatten()(inputs[0])
    x = Dense(1024, activation='relu')(x)
    x = BatchNormalization()(x)
    x = Dropout(0.5)(x)
    x = Dense(512, activation='relu')(x)
    x = BatchNormalization()(x)
    combined = Concatenate()([x] + aux_feats)
    x = Dense(256, activation='relu')(combined)
    x = Dense(128, activation='relu')(x)
    predictions = Dense(num_classes, activation='softmax')(x)
    model = Model(inputs=inputs, outputs=predictions)
    model.compile(optimizer=Adam(lr=0.0005), loss='categorical_crossentropy', metrics=['accuracy'])
    return model''')

    para("ResNet50 was next, similar idea but deeper base, pre trained imagenet, but I made it trainable and froze only the first layers, wait actually the last 50 trainable I think, no, the code says for layer in base model layers minus 50, set false, so early layers frozen. Global pool, concat aux, dense 512 relu, batch norm, dropout 0.5, dense 256, softmax. Adam with tiny lr 1e-5. It took longer and did not perform great, maybe because of the depth and not enough training time.", bold=True)
    code_block('''def build_resnet50_model(num_classes=30, input_shape=(224, 224, 3)):
    inputs, aux_feats = get_multi_input_layers(input_shape)
    base_model = ResNet50(weights='imagenet', include_top=False, input_tensor=inputs[0])
    base_model.trainable = True
    for layer in base_model.layers[:-50]:
        layer.trainable = False
    x = GlobalAveragePooling2D()(base_model.output)
    combined = Concatenate()([x] + aux_feats)
    x = Dense(512, activation='relu')(combined)
    x = BatchNormalization()(x)
    x = Dropout(0.5)(x)
    x = Dense(256, activation='relu')(x)
    predictions = Dense(num_classes, activation='softmax')(x)
    model = Model(inputs=inputs, outputs=predictions)
    model.compile(optimizer=Adam(lr=1e-5), loss='categorical_crossentropy', metrics=['accuracy'])
    return model''')

    para("For the SVM, its different, not a full model but extract features using the MobileNetV2 as feature extractor, predict on batches from train gen, up to 200 batches or whatever min len, stack the features which are like 1280 from the pool I guess plus aux making 1504 dims, labels from argmax. Then fit SVC with rbf kernel, C 10 gamma scale, probability true. Super fast, no epochs, just fit once.", bold=True)
    code_block('''def train_svm(train_gen, val_gen, feature_extractor, save_path):
    train_features, train_labels = [], []
    for i in range(min(len(train_gen), 200)):
        batch_inputs, batch_y = train_gen[i]
        features = feature_extractor.predict(batch_inputs, verbose=0)
        train_features.append(features)
        train_labels.append(np.argmax(batch_y, axis=1))

    train_features = np.vstack(train_features)
    train_labels = np.concatenate(train_labels)

    svm = SVC(kernel='rbf', C=10, gamma='scale', probability=True)
    svm.fit(train_features, train_labels)
    return svm''')

    para("Training for the deep ones used callbacks, model checkpoint on val acc save best, early stop patience 8 restore best, reduce lr on val loss factor 0.2 patience 4 min 1e-7. Fit with 20 epochs but it stopped early if no improve.", bold=True)
    code_block('''checkpoint = ModelCheckpoint(save_path, monitor='val_accuracy', save_best_only=True)
early_stop = EarlyStopping(monitor='val_accuracy', patience=8, restore_best_weights=True)
reduce_lr = ReduceLROnPlateau(monitor='val_loss', factor=0.2, patience=4, min_lr=1e-7)

history = model.fit(
    train_gen, epochs=20,
    validation_data=val_gen,
    callbacks=[checkpoint, early_stop, reduce_lr]
)''')
    doc.add_page_break()

    # =========================================================
    # 8 PREDICTION RESULTS
    # =========================================================
    heading("8 Prediction Results with Accuracy Table and Graphs", 1)

    para("Now the results, in the accuracy table, CNN MobileNet got 90.1 train 85.0 val, loss 0.4344, 9 epochs 5580 seconds. SVM 84.2 82.3 0.0000 1 epoch 120 s. ANN 43.0 39.7 1.7800 5 epochs 320 s. ResNet 29.1 25.1 3.1200 5 epochs 950 s. So CNN did best overall. The graph for train vs val acc shows some overfitting, like curves diverging a bit in ANN and ResNet especially. Training curves, accuracy rising loss dropping, steady for CNN but ANN loss stayed high.")

    sorted_data = sorted(MODEL_DATA, key=lambda x: x['val_accuracy'], reverse=True) if MODEL_DATA else []
    if sorted_data:
        add_table(
            ["Model", "Train Acc (%)", "Val Acc (%)", "Val Loss", "Epochs", "Time (s)"],
            [[d['model_name'], f"{d['train_accuracy']*100:.1f}", f"{d['val_accuracy']*100:.1f}",
              f"{d['val_loss']:.4f}", str(d['epochs_trained']), f"{d['training_time_seconds']:.0f}"]
             for d in sorted_data]
        )
    else:
        add_table(
            ["Model", "Train Acc", "Val Acc", "Val Loss", "Epochs", "Time"],
            [["CNN MobileNetV2","90.1","85.0","0.4344","9","5580"],
             ["SVM","84.2","82.3","0.0000","1","120"],
             ["ANN MLP","43.0","39.7","1.7800","5","320"],
             ["ResNet50","29.1","25.1","3.1200","5","950"]]
        )
    add_img(f"{RESULTS_DIR}/accuracy_table.png", 5.5, "Figure 8.1 Model Performance Summary")

    add_img(f"{RESULTS_DIR}/model_comparison.png", 5.5, "Figure 8.2 Train vs Validation Accuracy for all models")

    add_img(f"{RESULTS_DIR}/training_curves_all.png", 6.0, "Figure 8.3 Training curves accuracy and loss for all four models")

    para("Sample predictions, like tomato leaf 01 predicted late blight 92.3 diseased, apple fresh 02 apple fresh 88.7 healthy, grape rotten 03 grape rotten 85.1 diseased, pepper spot 04 pepper bell bacterial spot 91.5 diseased, banana fresh 05 banana fresh 87.2 healthy. Confidences over 85 mostly, seems decent.")
    add_table(
        ["Input Image", "Predicted Disease", "Confidence", "Status"],
        [["tomato_leaf_01.jpg", "Tomato Late Blight", "92.3%", "Diseased"],
         ["apple_fresh_02.jpg", "Apple Fresh", "88.7%", "Healthy"],
         ["grape_rotten_03.jpg", "Grape Rotten", "85.1%", "Diseased"],
         ["pepper_spot_04.jpg", "Pepper Bell Bacterial Spot", "91.5%", "Diseased"],
         ["banana_fresh_05.jpg", "Banana Fresh", "87.2%", "Healthy"]]
    )
    if os.path.exists(f"{RESULTS_DIR}/prediction_examples.png"):
        add_img(f"{RESULTS_DIR}/prediction_examples.png", 5.5, "Figure 8.4 Sample prediction results")
    doc.add_page_break()

    # =========================================================
    # 9 COMPARISON
    # =========================================================
    heading("9 Comparison Proposed Work vs Existing Work", 1)
    para("Comparing to other work, I looked at papers but our dataset is mixed fruits and leaves, harder than just PlantVillage. Table, AlexNet Mohanty 2016 PlantVillage 38 classes 99.3 val acc, VGG16 Too 2019 same 99. Graph shows ours lower, like 85 vs their 99. The differences, I think its because they used one clean dataset, consistent images, ours combined two with different lights backgrounds quality. Plus our classes include subtle fruit stuff like fresh rotten formalin, not just obvious leaf spots. We did only 20 epochs early stop, maybe more tuning helps. MobileNet is small 3.4M params for speed, not like VGG 138M.")

    add_table(
        ["Method", "Dataset", "Classes", "Val Accuracy (%)"],
        [["AlexNet (Mohanty et al 2016)", "PlantVillage", "38", "99.3"],
         ["VGG16 (Too et al 2019)", "PlantVillage", "38", "99.5"],
         ["InceptionV3 (Brahimi et al 2017)", "PlantVillage", "38", "99.2"],
         ["ResNet50 (Ferentinos 2018)", "PlantVillage", "38", "99.5"],
         ["DenseNet121 (Too et al 2019)", "PlantVillage", "38", "99.8"],
         ["Our CNN (MobileNetV2)", "PV + Fruits", "30", "85.1"],
         ["Our SVM", "PV + Fruits", "30", "82.3"]]
    )

    add_img(f"{RESULTS_DIR}/comparison_with_literature.png", 6.0, "Figure 9.1 Accuracy comparison published methods vs our approach")

    doc.add_page_break()

    # =========================================================
    # 10 RESULTS AND DISCUSSION
    # =========================================================
    heading("10 Results and Discussion", 1)
    para("In discussion, the results show trade offs, CNN MobileNet best at 85 val acc, transfer learning helps with imagenet features, frozen base plus aux sakaguchi tensors balance things. SVM close at 82, fast training 120 s vs 5580, handles high dim well with rbf. ANN bad 39, yeah flattening raw pixels no good without convs. ResNet lowest 25, surprising but too many params 25M trainable, low lr 1e-5, only 5 epochs, could do better with more time. The sakaguchi adds 224 features from those scales, multi scale context, but hard to say exactly without ablation.")
    para("The aux features, from dense on pooled small inputs, provide some spatial info at different resolutions, maybe helps a bit. Overall, it seems transfer learning wins here.")
    doc.add_page_break()

    # =========================================================
    # 11 CONCLUSIONS WITH FUTURE ENHANCEMENT
    # =========================================================
    heading("11 Conclusions with Future Enhancement", 1)

    para("For conclusions, the project made a system with multi inputs models for 30 classes, MobileNet best 85 acc reasonable time, SVM alternative fast 82 in 2 min, ANN not suited without convs, mixing datasets harder. Streamlit interface works for web use.")

    para("Future stuff, add real field images for better robustness, ensemble CNN SVM weighted vote for tough cases, android app with tflite offline, severity grading mild moderate severe for treatment, attention like grad cam to show what model sees, trust building. IoT integration alerts from cameras, and ablation on sakaguchi each scale contribution.")
    doc.add_page_break()

    # =========================================================
    # 12 REFERENCES
    # =========================================================
    heading("12 References", 1)
    refs = [
        "[1] S. P. Mohanty, D. P. Hughes, and M. Salathe, \"Using deep learning for image-based plant disease detection,\" Frontiers in Plant Science, vol. 7, p. 1419, 2016.",
        "[2] E. C. Too, L. Yujian, S. Njuki, and L. Yingchun, \"A comparative study of fine-tuning deep learning models for plant disease identification,\" Computers and Electronics in Agriculture, vol. 161, pp. 272-279, 2019.",
        "[3] M. Brahimi, K. Boukhalfa, and A. Moussaoui, \"Deep learning for tomato diseases: Classification and symptoms visualization,\" Applied Artificial Intelligence, vol. 31, no. 4, pp. 299-315, 2017.",
        "[4] K. P. Ferentinos, \"Deep learning models for plant disease detection and diagnosis,\" Computers and Electronics in Agriculture, vol. 145, pp. 311-318, 2018.",
        "[5] M. Sandler, A. Howard, M. Zhu, A. Zhmoginov, and L.-C. Chen, \"MobileNetV2: Inverted residuals and linear bottlenecks,\" in Proceedings of CVPR, pp. 4510-4520, 2018.",
        "[6] K. He, X. Zhang, S. Ren, and J. Sun, \"Deep residual learning for image recognition,\" in Proceedings of CVPR, pp. 770-778, 2016.",
        "[7] PlantVillage Dataset, Kaggle, https://www.kaggle.com/datasets/emmarex/plantdisease",
        "[8] Fruits Disease Dataset, Kaggle, https://www.kaggle.com/datasets/sriramr/fruits-diseases",
        "Project Repository: https://github.com/Azaan-5025/Fruit_Leaf_Disease_Detection"
    ]
    for ref in refs:
        para(ref)
    
    # Save
    doc.save(OUTFILE)
    print(f"\n✅ Report saved: {OUTFILE}")



if __name__ == "__main__":
    print("Generating Project Report...")
    ensure_dirs()
    build_report()
